import os
import subprocess
import time
import shutil
import pandas as pd
from boto3.session import Session
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from openai import OpenAI
from pymongo import MongoClient
from tqdm import tqdm
from typing import Dict
import uuid
import threading
from manim import *
from IPython.display import Video
from PIL import Image
from docx.shared import Pt
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from bson import ObjectId
from fastapi.middleware.cors import CORSMiddleware

ALLOWED_ORIGINS = [
    'http://localhost:3000/',
]

# MongoDB Connection
client = MongoClient('mongodb://localhost:27017/')
db = client.education_db

# MongoDB Collections
books_col = db.books
grades_col = db.grades
chapters_col = db.chapters
headings_col = db.headings
topics_col = db.topics
quizzes_col = db.quizzes
tasks_col = db.tasks

app = FastAPI()
# app.add_middleware(OriginMiddleware)  # Add the middleware to the FastAPI app
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)



# Create a Boto3 session
session = Session(aws_access_key_id=AWS_ACCESS_KEY_ID,
                  aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
                  region_name=REGION_NAME)

s3_client = session.client('s3')

openai_api_key = 'sk-'
os.environ['OPENAI_API_KEY'] = openai_api_key

openai_client = OpenAI(api_key=openai_api_key)


def extract_values(csv_file):
    data = pd.read_csv(csv_file)
    DATA = {}
    for _, row in data.iterrows():
        book = row['Book']
        grade = row['Grade']
        chapter = row['Chapters']
        sub_heading = row['Sub Chapters']
        topic = row['Topics']

        if book not in DATA:
            DATA[book] = {}

        if grade not in DATA[book]:
            DATA[book][grade] = {}

        if chapter not in DATA[book][grade]:
            DATA[book][grade][chapter] = {}

        if sub_heading not in DATA[book][grade][chapter]:
            DATA[book][grade][chapter][sub_heading] = []

        DATA[book][grade][chapter][sub_heading].append(topic)

    return DATA


def store_to_mongodb(book, grade, chapter, heading, topic, content, quiz, video_link):  # Include video_link here
    # Upsert book and get its ID
    book_result = books_col.update_one({'name': book}, {'$setOnInsert': {'name': book}}, upsert=True)
    book_id = book_result.upserted_id or books_col.find_one({'name': book})['_id']

    # Upsert grade and get its ID
    grade_result = grades_col.update_one({'name': grade, 'book_id': book_id},
                                         {'$setOnInsert': {'name': grade, 'book_id': book_id}}, upsert=True)
    grade_id = grade_result.upserted_id or grades_col.find_one({'name': grade, 'book_id': book_id})['_id']

    # Upsert chapter and get its ID
    chapter_result = chapters_col.update_one({'name': chapter, 'grade_id': grade_id},
                                             {'$setOnInsert': {'name': chapter, 'grade_id': grade_id}}, upsert=True)
    chapter_id = chapter_result.upserted_id or chapters_col.find_one({'name': chapter, 'grade_id': grade_id})['_id']

    # Upsert heading and get its ID
    heading_result = headings_col.update_one({'name': heading, 'chapter_id': chapter_id},
                                             {'$setOnInsert': {'name': heading, 'chapter_id': chapter_id}}, upsert=True)
    heading_id = heading_result.upserted_id or headings_col.find_one({'name': heading, 'chapter_id': chapter_id})['_id']

    # Upsert topic and get its ID, including quiz and video_link
    topics_col.update_one(
        {'name': topic, 'heading_id': heading_id},
        {'$setOnInsert': {'name': topic, 'heading_id': heading_id, 'content': content, 'quiz': quiz,
                          'video_link': video_link}},
        upsert=True
    )


def retrieve_data_from_mongodb():
    all_data = {}
    books = books_col.find()
    for book in books:
        book_name = book['name']
        book_id = book['_id']
        all_data[book_name] = {}

        grades = grades_col.find({'book_id': book_id})
        for grade in grades:
            grade_name = grade['name']
            grade_id = grade['_id']
            all_data[book_name][grade_name] = {}

            chapters = chapters_col.find({'grade_id': grade_id})
            for chapter in chapters:
                chapter_name = chapter['name']
                chapter_id = chapter['_id']
                all_data[book_name][grade_name][chapter_name] = {}

                headings = headings_col.find({'chapter_id': chapter_id})
                for heading in headings:
                    heading_name = heading['name']
                    heading_id = heading['_id']
                    all_data[book_name][grade_name][chapter_name][heading_name] = {}

                    topics = topics_col.find({'heading_id': heading_id})
                    for topic in topics:
                        topic_name = topic['name']
                        topic_id = topic['_id']
                        topic_content = topic.get('content', '')
                        topic_video_link = topic.get('video_link', '')
                        quiz_content = topic.get('quiz', '')  # Directly retrieve quiz from the topic document
                        all_data[book_name][grade_name][chapter_name][heading_name][topic_name] = {
                            'content': topic_content,
                            'quiz': quiz_content,
                            'video_link': topic_video_link
                        }

    return all_data


'''
@MANIM CODES
'''

GPT_SYSTEM_INSTRUCTIONS = """Write a complete Manim scripts for animations in Python. Ensure correct variable declaration and Must provide correct Latex that will be rendered without any error. Generate code only in a complete and detailed way, don`t write text explanations. Never explain code. Never add functions. Never add comments or  infinte loops. Never use other library than Manim/math. Only provide complete code block with correct Manim Syntax. Use variables with length of maximum 2 characters. At the end use 'self.play'.

```
from manim import *
from math import *

class GenScene(Scene):
    def construct(self):
        # Write here
```"""


def wrap_prompt(prompt: str) -> str:
    """
      Wraps the prompt in the GPT instructions
    """
    return f"Animation Request: {prompt}"


def extract_code(text: str) -> str:
    """
      Extracts the code from the text generated by GPT-3.5 from the ``` ``` blocks
    """
    pattern = re.compile(r"```(.*?)```", re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    else:
        return text


def extract_construct_code(code_str: str) -> str:
    """
      Extracts the code from the construct method
    """
    pattern = r"def construct\(self\):([\s\S]*)"
    match = re.search(pattern, code_str)
    if match:
        return match.group(1)
    else:
        return ""


def code_static_corrector(code_response: str) -> str:
    """
      Corrects some static errors in the code
      GPT only has information until 2021, so it ocasionally generates code
      that is not compatible with the latest version of Manim
    """
    # Replace ShowCreation with Create
    # Solution: https://www.reddit.com/r/manim/comments/qcosuj/nameerror_name_showcreation_is_not_defined/
    code_response = code_response.replace("ShowCreation", "Create")
    latex_special_chars = {"$": "\\$",
                           "%": "\\%",
                           "&": "\\&"}

    for char, escaped_char in latex_special_chars.items():
        code_response = code_response.replace(char, escaped_char)

    return code_response


def create_file_content(class_name, code_response):
    """
    Creates the content of the file to be written with a dynamic class name.
    """
    return f"""# Manim code generated with OpenAI GPT
# Command to generate animation: manim {class_name}.py {class_name} --format=mp4 --media_dir . --custom_folders video_dir

from manim import *
from math import *

class {class_name}(Scene):
    def construct(self):
{code_static_corrector(code_response)}"""


def generate_code(prompt):
    """
    Generates code using OpenAI based on a given prompt.

    Args:
        prompt: The text prompt to provide to OpenAI.

    Returns:
        The generated code as a string.
    """

    response = openai_client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": GPT_SYSTEM_INSTRUCTIONS},
            {"role": "user", "content": wrap_prompt(prompt)}
        ]
    )

    code_response = extract_code(response.choices[0].message.content)
    code_response = extract_construct_code(code_response)

    return code_response


def upload_to_s3(video_file, book, grade, chapter, heading, topic):
    # Ensure all components are strings and remove any leading/trailing spaces
    book_str = str(book).strip()
    grade_str = str(grade).strip()
    chapter_str = str(chapter).strip()
    heading_str = str(heading).strip()
    topic_str = str(topic).strip()

    object_name = f"{book_str}/{grade_str}/{chapter_str}/{heading_str}/{topic_str}/{os.path.basename(video_file)}".replace(
        ',', '')
    object_name = object_name.replace(' ', '-')
    print('Object name:', object_name)

    # Check if the video file exists
    if not os.path.exists(video_file):
        print(f"Video file not found: {video_file}")
        return None

    try:
        s3_client.upload_file(video_file, BUCKET_NAME, object_name)
        video_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{object_name}"
        return video_url
    except Exception as e:
        print(f"Error uploading {video_file} to S3:", e)
        return None


def create_and_save_animation(prompt, book, grade, chapter, heading, topic):
    generated_code = generate_code(prompt)
    print("Generated code:\n")
    print(generated_code)

    # Generate a unique class name using timestamp
    timestamp = time.strftime("%Y%m%d%H%M%S")
    unique_class_name = f"Scene_{timestamp}"
    video_file = f"{unique_class_name}.mp4"
    py_file = f"{unique_class_name}.py"

    try:
        with open(f"{unique_class_name}.py", "w") as f:
            f.write(create_file_content(unique_class_name, generated_code))

        py_link = upload_to_s3(py_file, book, grade, chapter, heading, topic)
        command_to_render = f"manim '{unique_class_name}.py' {unique_class_name} --format=mp4 --media_dir . --custom_folders"
        subprocess.run(command_to_render, check=True, shell=True)

        video_link = upload_to_s3(video_file, book, grade, chapter, heading, topic)

        return video_link
    except Exception as e:
        print("Error in creating or uploading animation:", e)
        return None


'''
CONTENT GENERATION PIPELINE
'''


def generate_content_and_quiz(result, grade):
    # Enhanced prompt for content generation
    content_prompt = f'''
    You are an experienced educator and content creator, specializing in crafting educational materials for Grade {grade} students. Your task is to generate clear, concise, and engaging explanations for the following topics. Incorporate diagrams, examples, and analogies where helpful to enhance understanding. Each explanation should be limited to 1-2 pages, making it suitable for a Grade {grade} student.

    Topics:
    {result}
    '''

    # Generate educational content
    content_response = openai_client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": content_prompt},
            {"role": "user", "content": ""}
        ]
    )
    content = content_response.choices[0].message.content

    # Enhanced prompt for quiz generation
    quiz_prompt = f'''
    You are an expert in creating educational quizzes for Grade {grade} students. Based on the content provided, craft a set of quiz questions that accurately assess the understanding of the topics. Ensure the questions are clear, concise, and suitable for Grade {grade} students, Quiz should be of questions/answering type and don`t provide solution just 5 questions.

    Content:
    {content}
    '''

    # Generate quiz
    quiz_response = openai_client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": quiz_prompt},
            {"role": "user", "content": ""}
        ]
    )
    quiz = quiz_response.choices[0].message.content

    return content, quiz


def upload_docx_to_s3(file_name, bucket, object_name=None):
    if object_name is None:
        object_name = str(file_name)
    try:
        response = s3_client.upload_file(str(file_name), bucket, object_name)
        return f"https://{bucket}.s3.amazonaws.com/{object_name}"
    except Exception as e:
        print(f"Error uploading {file_name} to S3:", e)
        return None


from bson import ObjectId


def regenerate_content_quiz_video(topic_id_str):
    # Convert the string ID to ObjectId
    try:
        topic_id = ObjectId(topic_id_str)
    except Exception as e:
        print(f"Invalid ID format: {e}")
        return

    # Retrieve the topic from MongoDB
    topic_data = topics_col.find_one({'_id': topic_id})
    if not topic_data:
        print("Topic not found.")
        return

    # Retrieve the heading
    heading_id = topic_data.get('heading_id')
    heading_data = headings_col.find_one({'_id': heading_id})
    if not heading_data:
        print("Heading not found.")
        return
    heading_name = heading_data.get('name')

    # Retrieve the chapter
    chapter_id = heading_data.get('chapter_id')
    chapter_data = chapters_col.find_one({'_id': chapter_id})
    if not chapter_data:
        print("Chapter not found.")
        return
    chapter_name = chapter_data.get('name')

    # Retrieve the grade
    grade_id = chapter_data.get('grade_id')
    grade_data = grades_col.find_one({'_id': grade_id})
    if not grade_data:
        print("Grade not found.")
        return
    grade_name = grade_data.get('name')

    # Retrieve the book
    book_id = grade_data.get('book_id')
    book_data = books_col.find_one({'_id': book_id})
    if not book_data:
        print("Book not found.")
        return
    book_name = book_data.get('name')

    topic_name = topic_data.get('name')

    print(f'[INFO] Generating New Contents for grade {grade_name} , {topic_name}...')

    # Generate new content and quiz
    new_content, new_quiz = generate_content_and_quiz(topic_name, grade_name)

    # Update MongoDB with new content and quiz
    topics_col.update_one(
        {'_id': topic_id},
        {'$set': {'content': new_content, 'quiz': new_quiz}}
    )

    # Generate a new video
    video_link = create_and_save_animation(new_content, book_name, grade_name, chapter_name, heading_name, topic_name)

    # Update MongoDB with new video link
    if video_link:
        topics_col.update_one(
            {'_id': topic_id},
            {'$set': {'video_link': video_link}}
        )
    return topic_id


import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def create_enhanced_docx(all_data):
    document = docx.Document()

    for book, grades in all_data.items():
        print(f"Adding book: {book}")  # Debugging
        add_heading(document, book, 0)  # Book name as the main heading

        for grade, chapters in grades.items():
            grade_heading = f"Grade {grade}"
            print(f"Adding grade: {grade_heading}")  # Debugging
            add_heading(document, grade_heading, 1)

            for chapter, headings in chapters.items():
                print(f"Adding chapter: {chapter}")  # Debugging
                add_heading(document, chapter, 2)

                for heading, topics in headings.items():
                    print(f"Adding heading: {heading}")  # Debugging
                    add_heading(document, heading, 3)

                    for topic, details in topics.items():
                        print(f"Adding topic: {topic}")  # Debugging
                        add_heading(document, topic, 4)
                        add_paragraph(document, details['content'])  # Topic content
                        add_heading(document, 'Quiz', 4)
                        add_paragraph(document, details['quiz'])  # Quiz content

    filename = 'formatted_document.docx'
    document.save(filename)
    return filename


def upload_docx_to_s3(file_name, bucket, object_name=None):
    if object_name is None:
        object_name = str(file_name)
    try:
        response = s3_client.upload_file(str(file_name), bucket, object_name)
        return f"https://{bucket}.s3.amazonaws.com/{object_name}"
    except Exception as e:
        print(f"Error uploading {file_name} to S3:", e)
        return None


def add_heading(document, text, level):
    heading = document.add_heading(text, level)
    heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    heading_font = heading.runs[0].font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(16)  # Set font size to 12 for all headings
    heading_font.bold = True


def add_paragraph(document, text):
    cleaned_text = clean_text(text)
    paragraph = document.add_paragraph(cleaned_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if paragraph.runs:
        paragraph_font = paragraph.runs[0].font
        paragraph_font.name = 'Times New Roman'
        paragraph_font.size = Pt(12)  # Set font size to 12 for all paragraphs


def clean_text(text):
    text = re.sub(r"[^\w\s',!?:().*+-=]", "", text)
    text = re.sub(r"-{2,}", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.capitalize()  # Capitalize the first letter of the text


def retrieve_all_contents():
    data = retrieve_data_from_mongodb()
    return data


async def save_upload_file_tmp(upload_file: UploadFile, tmp_path: str = "/tmp") -> str:
    try:
        file_path = os.path.join(tmp_path, upload_file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
    finally:
        upload_file.file.close()
    return file_path


def process_csv_file(file: UploadFile):
    result = extract_values(csv_file=file)
    for book, grades in result.items():
        for grade, chapters in grades.items():
            for chapter, headings in chapters.items():
                for heading, topics in headings.items():
                    for topic in tqdm(topics):
                        print(
                            f"Processing: Book: {book}, Grade: {grade}, Chapter: {chapter}, Heading: {heading}, Topic: {topic}")
                        content, quiz = generate_content_and_quiz(topic, grade)
                        video_link = create_and_save_animation(content, book, grade, chapter, heading, topic)
                        if video_link:
                            store_to_mongodb(book, grade, chapter, heading, topic, content, quiz, video_link)
                        else:
                            store_to_mongodb(book, grade, chapter, heading, topic, content, quiz, None)

    data = retrieve_data_from_mongodb()
    filename = create_enhanced_docx(data)
    s3_bucket = "brainjee-model"
    s3_link = upload_docx_to_s3(filename, s3_bucket)

    if s3_link:
        print(f"Document uploaded successfully: {s3_link}")
        return s3_link
    else:
        print("Failed to upload the document to S3.")
        return None


# FastAPI Endpoints


def background_task(task_id: str, file_path: str):
    try:
        # Process the file
        process_csv_file(file_path)
        tasks_col.update_one(
            {'_id': ObjectId(task_id)},
            {'$set': {'status': 'Completed', 'result': "Success Message or URL"}}
        )
    except Exception as e:
        tasks_col.update_one(
            {'_id': ObjectId(task_id)},
            {'$set': {'status': 'Failed', 'error': str(e)}}
        )


@app.post("/generate-contents/")
async def generate_contents(upload_file: UploadFile = File(...)):
    task_id = ObjectId()
    tasks_col.insert_one({'_id': task_id, 'status': 'Processing'})

    file_path = await save_upload_file_tmp(upload_file)
    threading.Thread(target=background_task, args=(str(task_id), file_path)).start()

    return {"task_id": str(task_id)}


@app.get("/task-status/{task_id}")
async def get_task_status(task_id: str):
    task = tasks_col.find_one({'_id': ObjectId(task_id)})
    if task:
        return {"status": task.get('status'), "result": task.get('result'), "error": task.get('error')}
    else:
        raise HTTPException(status_code=404, detail="Task not found")


@app.post("/regenerate-content/{topic_id}")
async def regenerate_content_endpoint(topic_id: str):
    task_id = ObjectId()
    tasks_col.insert_one({'_id': task_id, 'status': 'Processing'})

    threading.Thread(target=background_regenerate_content, args=(str(task_id), topic_id)).start()
    return {"task_id": str(task_id)}


from bson import ObjectId


def background_regenerate_content(task_id: str, topic_id_str: str):
    try:
        topic_id = ObjectId(topic_id_str)
        regenerated_topic_id = regenerate_content_quiz_video(topic_id)

        if regenerated_topic_id:
            tasks_col.update_one(
                {'_id': ObjectId(task_id)},
                {'$set': {'status': 'Completed', 'result': str(regenerated_topic_id)}}
            )
        else:
            tasks_col.update_one(
                {'_id': ObjectId(task_id)},
                {'$set': {'status': 'Failed', 'error': "Content regeneration failed"}}
            )
    except Exception as e:
        tasks_col.update_one(
            {'_id': ObjectId(task_id)},
            {'$set': {'status': 'Failed', 'error': str(e)}}
        )


@app.get("/retrieve-contents/")
async def retrieve_contents():
    contents = retrieve_all_contents()
    return JSONResponse(content=contents)
